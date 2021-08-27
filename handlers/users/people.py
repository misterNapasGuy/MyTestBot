import datetime
from unittest.case import _id

from aiogram.dispatcher.filters import Text
from aiogram.types import Message, CallbackQuery
from docx import Document

from filters.states import From
from keyboards.default.default_menus import add_and_main
from keyboards.inline.inline_menus import get_inline_menu_groups, cb_id_group, get_inline_menu_users, cb_id_users, \
    get_admin_private_task_menu, get_admin_load_report_menu, cb_id_load_report_users
from loader import dp
from moduls.notifications import send_notification, get_admins_all
from moduls.personnel import get_groups, get_group_name, get_users, is_admin_check, get_user, get_tasks_and_users, \
    get_group_tasks_result, update_user_id_private_task, update_group_id_public_task
from moduls.tasks import get_tasks_privat, create_task


@dp.message_handler(Text(equals=["Кадры"]))
async def check_private_tasks(message: Message):
    await From.add_group.set()
    await message.answer(text="Меню кадров", reply_markup=add_and_main)
    await dp.bot.send_message(chat_id=message.chat.id, text="Ваши группы:",
                              reply_markup=get_inline_menu_groups(groups=get_groups()))
    await dp.bot.send_message(chat_id=message.chat.id, text="Отчет по сотрудникам и группам за сегодня:",
                              reply_markup=get_admin_load_report_menu(get_users(_id)))


@dp.callback_query_handler(cb_id_load_report_users.filter(), state="*")
async def load_report_users(call: CallbackQuery):
    report_string = get_tasks_and_users()
    res = ["Выполнено", "Не выполнено"]
    document = Document()
    d = datetime.date.today()
    document.add_heading(f'Отчет по работе сотрудников  за {d.day}.{d.month}.{d.year}')
    document.add_heading("Личные задания сотрудников", level=1)
    for i in report_string:
        document.add_paragraph(f"Имя:{i[2]}\nЗадание:{i[1]} - {res[0] if i[3] else res[1]}\n")
    document.add_heading("Групповые задания", level=1)
    report_group_string = get_group_tasks_result()
    for i in report_group_string:
        document.add_paragraph(f"Название группы:{i[1]}\nЗадание:{i[2]} - {res[0] if i[3] else res[1]}\n")
    document.save('Отчет.docx')
    with open('Отчет.docx', 'rb') as file:
        await dp.bot.send_document(chat_id=1064845609, document=file)
    update_group_id_public_task()
    update_user_id_private_task()
    await dp.bot.send_message(chat_id=1064845609, text='___Ваше фото обрабатывается..___', parse_mode="Markdown")


@dp.callback_query_handler(cb_id_group.filter(), state=From.add_group)
async def show_users(call: CallbackQuery, callback_data: dict):
    await From.add_user.set()
    await call.message.delete()
    _id = callback_data.get("id_group")
    await call.message.answer(text=f"Группа {get_group_name(_id)}:",
                              reply_markup=get_inline_menu_users(get_users(_id)))


@dp.callback_query_handler(cb_id_users.filter(), state=From.add_user)
async def show_users(call: CallbackQuery, callback_data: dict):
    await From.private_task_admin.set()
    await call.message.delete()
    await dp.bot.send_message(chat_id=call.from_user.id,
                              text=f"Задачи пользователя {get_user(callback_data.get('id_user'))[2]}")
    result = get_tasks_privat(id_user=callback_data.get("id_user"))
    if is_admin_check(call.from_user.id):
        for io in result:
            if io[5] is not None:
                await dp.bot.send_message(chat_id=call.from_user.id,
                                          text=f"{'✅ ' if io[3] else '⭕ '}{io[1]}\r\n📩 Коментарии:\r\n{io[5]}",
                                          reply_markup=get_admin_private_task_menu(io[0]))
            else:
                await dp.bot.send_message(chat_id=call.from_user.id, text=f"{'✅ ' if io[3] else '⭕ '}{io[1]}",
                                          reply_markup=get_admin_private_task_menu(io[0]))


@dp.message_handler(state=From.private_task_admin, text="Добавить")
async def add_private_task(message):
    await message.answer("Введите текст задачи")
    await From.private_task_admin2.set()


@dp.message_handler(state=From.private_task_admin2)
async def add_private_task2(message: Message):
    if message.text == "Редактировать" or message.text == "Добавить" or message.text == "↪ В главное меню":
        await message.answer(text="Данные не верны!")
        await From.add_group.set()
    create_task(content=f"{message.text}", user_id=message.from_user.id)
    await From.add_group.set()
    await send_notification(IDs=[i[0] for i in get_admins_all(id_user=message.from_user.id,
                                                              name_admins_group="admins")],
                            content=f"📰 {message.from_user.full_name} добавил себе задание: {message.text}")
    await message.answer("Задача добавлена!")
